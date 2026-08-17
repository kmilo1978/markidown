import os
import re
import io
import tempfile
import zipfile
import json
import xml.etree.ElementTree as ET
import requests as http_requests
from flask import Flask, render_template, request, jsonify, send_file
from markitdown import MarkItDown
from urllib.parse import urlparse
import markdown as md_lib
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024

ALLOWED_EXTENSIONS = {
    "pdf", "docx", "doc", "pptx", "ppt", "xlsx", "xls",
    "html", "htm", "txt", "csv", "json", "xml",
    "jpg", "jpeg", "png", "gif", "bmp", "tiff",
    "mp3", "wav", "m4a",
    "zip", "epub", "ipynb",
}

# Presets define frontmatter and cleanup rules per use case
PRESETS = {
    "general": {
        "tags": ["converted"],
        "clean_level": "basic",
        "description": "Conversion estandar sin filtros adicionales",
    },
    "seo": {
        "tags": ["seo", "content-audit"],
        "clean_level": "aggressive",
        "extract_meta": True,
        "description": "Optimizado para analisis SEO: extrae meta tags, limpia navegacion",
    },
    "documentation": {
        "tags": ["docs", "reference"],
        "clean_level": "moderate",
        "keep_code_blocks": True,
        "description": "Preserva bloques de codigo y estructura tecnica",
    },
    "research": {
        "tags": ["research", "source"],
        "clean_level": "moderate",
        "add_citation": True,
        "description": "Agrega citas y metadatos de fuente para investigacion",
    },
    "archival": {
        "tags": ["archive", "snapshot"],
        "clean_level": "minimal",
        "add_timestamp": True,
        "description": "Archivado completo con timestamp y minima limpieza",
    },
}

# Patterns to remove during content normalization
NOISE_PATTERNS = {
    "nav": [
        r'(?m)^\s*\[.*?\]\(#.*?\)\s*\|?\s*$',  # nav links
        r'(?m)^.*?(Skip to|Jump to|Go to).*?$',
        r'(?m)^.*?(Menu|Navigation|Breadcrumb).*?$',
    ],
    "footer": [
        r'(?m)^.*?(All rights reserved|Copyright|\u00a9).*?$',
        r'(?m)^.*?(Privacy Policy|Terms of Service|Cookie Policy).*?$',
        r'(?m)^.*?(Powered by|Built with).*?$',
    ],
    "cookies": [
        r'(?m)^.*?(cookie|Cookie|COOKIE).*?(accept|consent|policy|banner).*?$',
        r'(?m)^.*?(We use cookies|This site uses cookies).*?$',
    ],
    "ads": [
        r'(?m)^.*?(Advertisement|Sponsored|Ad\s*:).*?$',
        r'(?m)^.*?(Subscribe now|Sign up for|Newsletter).*?$',
    ],
}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def normalize_content(text, level="basic"):
    """Remove navigation, footers, ads, cookie banners based on clean level."""
    if level == "minimal":
        return text

    patterns_to_apply = []
    if level in ("basic", "moderate", "aggressive"):
        patterns_to_apply.extend(NOISE_PATTERNS["cookies"])
    if level in ("moderate", "aggressive"):
        patterns_to_apply.extend(NOISE_PATTERNS["nav"])
        patterns_to_apply.extend(NOISE_PATTERNS["footer"])
    if level == "aggressive":
        patterns_to_apply.extend(NOISE_PATTERNS["ads"])

    for pattern in patterns_to_apply:
        text = re.sub(pattern, '', text)

    # Remove excessive blank lines
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    return text.strip()


def should_compress(content, source_type="single"):
    """Determine if content should be compressed before sending to an agent.
    Returns recommendation dict with compress boolean, reason, and estimated savings."""
    content_length = len(content)
    token_estimate = content_length // 4  # rough estimate: 1 token ~ 4 chars
    
    # Thresholds
    SINGLE_THRESHOLD = 8000  # ~2000 tokens
    BULK_THRESHOLD = 3000    # Lower threshold per page in bulk mode
    
    recommendation = {
        "compress": False,
        "reason": "",
        "token_estimate": token_estimate,
        "estimated_savings": "0%",
        "content_type": "prose",
        "compressor": "kompress-v2-base",
    }
    
    # Detect content type for compressor routing
    if content.count("{") > 10 or content.count("[") > 10:
        recommendation["content_type"] = "structured_data"
        recommendation["compressor"] = "SmartCrusher"
        recommendation["estimated_savings"] = "60-95%"
    elif content.count("def ") > 3 or content.count("function ") > 3 or content.count("class ") > 3:
        recommendation["content_type"] = "code"
        recommendation["compressor"] = "CodeCompressor"
        recommendation["estimated_savings"] = "15-20%"
    else:
        recommendation["content_type"] = "prose"
        recommendation["compressor"] = "Kompress-v2-base"
        recommendation["estimated_savings"] = "30-50%"
    
    # Decision logic
    if source_type == "bulk":
        if content_length > BULK_THRESHOLD:
            recommendation["compress"] = True
            recommendation["reason"] = f"Bulk mode: {token_estimate} tokens estimados. Comprimir ahorra contexto para procesar multiples paginas."
        else:
            recommendation["reason"] = f"Contenido corto ({token_estimate} tokens). Compresion no necesaria."
    elif source_type == "single":
        if content_length > SINGLE_THRESHOLD:
            recommendation["compress"] = True
            recommendation["reason"] = f"Documento largo ({token_estimate} tokens). Recomendado comprimir antes de enviar al agente."
        else:
            recommendation["reason"] = f"Documento manejable ({token_estimate} tokens). Compresion opcional."
    
    return recommendation


def apply_preset(markdown_content, preset_name, source_url=None, filename=None):
    """Apply preset transformations and generate frontmatter."""
    preset = PRESETS.get(preset_name, PRESETS["general"])

    # Normalize content
    cleaned = normalize_content(markdown_content, preset["clean_level"])

    # Build frontmatter
    fm_lines = ["---"]
    fm_lines.append(f"preset: {preset_name}")
    fm_lines.append(f"converted: {__import__('datetime').date.today().isoformat()}")

    if source_url:
        fm_lines.append(f"source_url: {source_url}")
    if filename:
        fm_lines.append(f"original_file: {filename}")

    if preset.get("add_citation") and source_url:
        fm_lines.append(f"citation: \"Retrieved from {source_url} on {__import__('datetime').date.today().isoformat()}\"")

    if preset.get("add_timestamp"):
        fm_lines.append(f"archived_at: {__import__('datetime').datetime.now().isoformat()}")

    fm_lines.append("tags:")
    for tag in preset["tags"]:
        fm_lines.append(f"  - {tag}")

    fm_lines.append("---")
    fm_lines.append("")

    return "\n".join(fm_lines) + cleaned


def fetch_url(url):
    """Download content from URL, return (bytes, extension, netloc)."""
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError("La URL debe comenzar con http:// o https://")

    response = http_requests.get(url, timeout=30, headers={
        "User-Agent": "Mozilla/5.0 (compatible; MarkItDown/1.0)"
    })
    response.raise_for_status()

    path = parsed.path
    ext = os.path.splitext(path)[1].lower() if "." in path else ""
    if not ext:
        content_type = response.headers.get("Content-Type", "")
        if "pdf" in content_type:
            ext = ".pdf"
        elif "word" in content_type or "docx" in content_type:
            ext = ".docx"
        else:
            ext = ".html"

    return response.content, ext, parsed.netloc


def markdown_to_html(md_text):
    """Convert markdown to styled HTML."""
    html_body = md_lib.markdown(md_text, extensions=['tables', 'fenced_code', 'toc'])
    html_doc = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; max-width: 800px; margin: 2rem auto; padding: 0 1rem; line-height: 1.6; color: #333; }}
h1, h2, h3 {{ margin-top: 1.5em; }}
code {{ background: #f4f4f4; padding: 0.2em 0.4em; border-radius: 3px; font-size: 0.9em; }}
pre {{ background: #f4f4f4; padding: 1rem; border-radius: 6px; overflow-x: auto; }}
pre code {{ background: none; padding: 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
th, td {{ border: 1px solid #ddd; padding: 0.5rem; text-align: left; }}
th {{ background: #f4f4f4; }}
blockquote {{ border-left: 4px solid #ddd; margin: 1rem 0; padding: 0.5rem 1rem; color: #666; }}
img {{ max-width: 100%; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
    return html_doc


def markdown_to_docx(md_text):
    """Convert markdown to DOCX document."""
    doc = Document()
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        # List items
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
        # Blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'
        # Regular paragraph
        elif line.strip():
            doc.add_paragraph(line)
        # Empty line = skip
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/presets", methods=["GET"])
def get_presets():
    """Return available presets."""
    return jsonify(PRESETS)


@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "No se envio ningun archivo"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No se selecciono ningun archivo"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": f"Tipo de archivo no soportado. Extensiones permitidas: {', '.join(sorted(ALLOWED_EXTENSIONS))}"}), 400

    preset_name = request.form.get("preset", "general")

    suffix = os.path.splitext(file.filename)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        file.save(tmp.name)
        tmp_path = tmp.name

    try:
        md = MarkItDown()
        result = md.convert(tmp_path)
        markdown_content = apply_preset(result.text_content, preset_name, filename=file.filename)
    except Exception as e:
        return jsonify({"error": f"Error al convertir: {str(e)}"}), 500
    finally:
        os.unlink(tmp_path)

    return jsonify({
        "filename": file.filename,
        "markdown": markdown_content,
        "preset": preset_name,
    })


@app.route("/convert-url", methods=["POST"])
def convert_url():
    data = request.get_json()
    if not data or "url" not in data:
        return jsonify({"error": "No se proporciono una URL"}), 400

    url = data["url"].strip()
    preset_name = data.get("preset", "general")

    if not url:
        return jsonify({"error": "La URL esta vacia"}), 400

    try:
        content, ext, netloc = fetch_url(url)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except http_requests.exceptions.Timeout:
        return jsonify({"error": "Timeout: la URL tardo demasiado en responder"}), 400
    except http_requests.exceptions.RequestException as e:
        return jsonify({"error": f"Error al descargar la URL: {str(e)}"}), 400

    filename = netloc.replace(".", "_") + ext

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        md = MarkItDown()
        result = md.convert(tmp_path)
        markdown_content = apply_preset(result.text_content, preset_name, source_url=url)
    except Exception as e:
        return jsonify({"error": f"Error al convertir: {str(e)}"}), 500
    finally:
        os.unlink(tmp_path)

    return jsonify({
        "filename": filename,
        "url": url,
        "markdown": markdown_content,
        "preset": preset_name,
    })


@app.route("/convert-sitemap", methods=["POST"])
def convert_sitemap():
    data = request.get_json()
    if not data or "sitemap_url" not in data:
        return jsonify({"error": "No se proporciono la URL del sitemap"}), 400

    sitemap_url = data["sitemap_url"].strip()
    max_pages = data.get("max_pages", 20)
    preset_name = data.get("preset", "general")

    if not sitemap_url:
        return jsonify({"error": "La URL del sitemap esta vacia"}), 400

    try:
        response = http_requests.get(sitemap_url, timeout=30, headers={
            "User-Agent": "Mozilla/5.0 (compatible; MarkItDown/1.0)"
        })
        response.raise_for_status()
    except http_requests.exceptions.RequestException as e:
        return jsonify({"error": f"Error al descargar el sitemap: {str(e)}"}), 400

    try:
        root = ET.fromstring(response.content)
        ns = {"sm": "http://www.sitemaps.org/schemas/sitemap/0.9"}
        urls = []

        for url_elem in root.findall(".//sm:url/sm:loc", ns):
            urls.append(url_elem.text.strip())

        if not urls:
            for url_elem in root.iter():
                if url_elem.tag.endswith("loc"):
                    if url_elem.text:
                        urls.append(url_elem.text.strip())

    except ET.ParseError:
        return jsonify({"error": "El contenido no es un sitemap XML valido"}), 400

    if not urls:
        return jsonify({"error": "No se encontraron URLs en el sitemap"}), 400

    urls = urls[:max_pages]

    results = []
    errors = []
    md = MarkItDown()

    for page_url in urls:
        try:
            content, ext, netloc = fetch_url(page_url)

            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            try:
                convert_result = md.convert(tmp_path)
                parsed = urlparse(page_url)
                path_name = parsed.path.strip("/").replace("/", "_") or "index"
                filename = f"{netloc}_{path_name}.md"

                markdown_content = apply_preset(
                    convert_result.text_content, preset_name, source_url=page_url
                )

                results.append({
                    "url": page_url,
                    "filename": filename,
                    "markdown": markdown_content,
                })
            finally:
                os.unlink(tmp_path)

        except Exception as e:
            errors.append({"url": page_url, "error": str(e)})

    return jsonify({
        "total_urls": len(urls),
        "converted": len(results),
        "failed": len(errors),
        "results": results,
        "errors": errors,
        "preset": preset_name,
    })


@app.route("/export", methods=["POST"])
def export_markdown():
    """Reverse conversion: Markdown to HTML or DOCX."""
    data = request.get_json()
    if not data or "markdown" not in data:
        return jsonify({"error": "No se proporciono contenido markdown"}), 400

    md_text = data["markdown"]
    export_format = data.get("format", "html")
    filename = data.get("filename", "document")

    # Strip frontmatter before export
    if md_text.startswith("---"):
        end = md_text.find("---", 3)
        if end != -1:
            md_text = md_text[end + 3:].strip()

    if export_format == "html":
        html_content = markdown_to_html(md_text)
        buffer = io.BytesIO(html_content.encode('utf-8'))
        buffer.seek(0)
        return send_file(buffer, mimetype='text/html', as_attachment=True,
                        download_name=f"{filename}.html")

    elif export_format == "docx":
        buffer = markdown_to_docx(md_text)
        return send_file(buffer,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True, download_name=f"{filename}.docx")

    elif export_format == "pdf":
        try:
            from weasyprint import HTML as WeasyHTML
            html_content = markdown_to_html(md_text)
            buffer = io.BytesIO()
            WeasyHTML(string=html_content).write_pdf(buffer)
            buffer.seek(0)
            return send_file(buffer, mimetype='application/pdf', as_attachment=True,
                            download_name=f"{filename}.pdf")
        except ImportError:
            return jsonify({"error": "PDF export no disponible en este servidor (requiere weasyprint)"}), 501
        except Exception as e:
            return jsonify({"error": f"Error generando PDF: {str(e)}"}), 500

    return jsonify({"error": f"Formato no soportado: {export_format}"}), 400


@app.route("/compression-check", methods=["POST"])
def compression_check():
    """Check if content should be compressed before sending to an agent."""
    data = request.get_json()
    if not data or "markdown" not in data:
        return jsonify({"error": "No se proporciono contenido markdown"}), 400
    
    markdown = data["markdown"]
    source_type = data.get("source_type", "single")
    
    recommendation = should_compress(markdown, source_type)
    return jsonify(recommendation)


@app.route("/compression-bulk-check", methods=["POST"])
def compression_bulk_check():
    """Check compression recommendations for bulk sitemap results."""
    data = request.get_json()
    if not data or "results" not in data:
        return jsonify({"error": "No se proporcionaron resultados"}), 400
    
    results = data["results"]
    total_tokens = 0
    recommendations = []
    
    for item in results:
        rec = should_compress(item.get("markdown", ""), source_type="bulk")
        rec["url"] = item.get("url", "")
        rec["filename"] = item.get("filename", "")
        total_tokens += rec["token_estimate"]
        recommendations.append(rec)
    
    compress_count = sum(1 for r in recommendations if r["compress"])
    
    summary = {
        "total_pages": len(results),
        "total_tokens_estimate": total_tokens,
        "pages_to_compress": compress_count,
        "overall_recommendation": "compress" if total_tokens > 10000 else "optional",
        "reason": f"{total_tokens} tokens totales en {len(results)} paginas. "
                  + ("Recomendado usar Headroom para reducir contexto antes de enviar al agente." if total_tokens > 10000
                     else "Volumen manejable. Compresion opcional."),
        "headroom_command": "headroom proxy --port 8787" if total_tokens > 10000 else None,
        "per_page": recommendations,
    }
    
    return jsonify(summary)


@app.route("/compress", methods=["POST"])
def compress_content():
    """Compress markdown content using basic Headroom-style compression.
    Removes redundant URLs, deduplicates repeated patterns, and strips noise."""
    data = request.get_json()
    if not data or "markdown" not in data:
        return jsonify({"error": "No se proporciono contenido markdown"}), 400

    md_text = data["markdown"]
    original_length = len(md_text)
    original_tokens = original_length // 4

    # --- Compression strategies ---
    compressed = md_text

    # 1. Deduplicate repeated URLs (keep first, replace rest with [ref:N])
    import re as compress_re
    url_pattern = r'https?://[^\s\)\]\"\'<>]+'
    found_urls = compress_re.findall(url_pattern, compressed)
    url_counts = {}
    for u in found_urls:
        url_counts[u] = url_counts.get(u, 0) + 1

    url_refs = {}
    ref_index = 1
    for url, count in url_counts.items():
        if count > 1 and len(url) > 40:
            url_refs[url] = f"[ref:{ref_index}]"
            ref_index += 1

    # Build reference table
    ref_table = ""
    if url_refs:
        ref_table = "\n\n<!-- URL References -->\n"
        for url, ref in url_refs.items():
            ref_table += f"<!-- {ref} = {url} -->\n"
            # Replace all occurrences after the first
            first_pos = compressed.find(url)
            if first_pos >= 0:
                compressed = compressed[:first_pos + len(url)] + compressed[first_pos + len(url):].replace(url, ref)

    # 2. Remove redundant whitespace and blank lines
    compressed = compress_re.sub(r'\n{3,}', '\n\n', compressed)

    # 3. Remove tracking parameters from URLs
    compressed = compress_re.sub(r'[?&](utm_\w+|fbclid|gclid|ref|source|medium|campaign)=[^&\s\)\]]*', '', compressed)

    # 4. Collapse repeated markdown image badges (common in Product Hunt, shields.io)
    compressed = compress_re.sub(r'(\[!\[.*?\]\(.*?\)\]\(.*?\))\s*(?=\[!\[)', r'\1 ', compressed)

    # 5. Remove empty markdown links
    compressed = compress_re.sub(r'\[]\(.*?\)', '', compressed)

    # 6. Strip excessive badge/shield patterns
    compressed = compress_re.sub(r'!\[.*?\]\(https://img\.shields\.io/.*?\)', '[badge]', compressed)

    # Add reference table at end
    compressed = compressed + ref_table

    compressed_length = len(compressed)
    compressed_tokens = compressed_length // 4
    savings_pct = round((1 - compressed_length / max(original_length, 1)) * 100, 1)

    return jsonify({
        "original_tokens": original_tokens,
        "compressed_tokens": compressed_tokens,
        "savings_percent": savings_pct,
        "compressed_markdown": compressed,
    })


if __name__ == "__main__":
    app.run(debug=True, port=5000)
